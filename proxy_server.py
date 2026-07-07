# -*- coding: utf-8 -*-
from __future__ import annotations

import gzip
import http.server
import io
import json
import mimetypes
import os
import socket as _socket
import ssl
import threading
import time as _time
import traceback
import urllib.error
import urllib.parse
import urllib.request
from pathlib import Path

# 懒加载 datahub（避免 import 阶段死锁）
_datahub = None
_datahub_ok = False
_datahub_lock = _time.time()  # 记录首次尝试时间
_scheduler_started = False

def _get_datahub():
    """获取 datahub 服务实例，超时则返回 None"""
    global _datahub, _datahub_ok
    if _datahub is not None:
        return _datahub if _datahub_ok else None
    # 超过30秒还没导入成功就放弃
    if _time.time() - _datahub_lock > 30:
        return None
    try:
        from backend.datahub import datahub_service
        _datahub = datahub_service
        _datahub_ok = True
        return _datahub
    except Exception:
        _datahub = "failed"
        return None

def _get_scheduler_funcs():
    try:
        from backend.tasks.scheduler import start_scheduler, stop_scheduler
        return start_scheduler, stop_scheduler
    except Exception:
        return None, None


def _start_scheduler_once():
    global _scheduler_started
    if _scheduler_started:
        return True
    try:
        start_scheduler, _ = _get_scheduler_funcs()
        if start_scheduler is None:
            return False
        start_scheduler()
        _scheduler_started = True
        return True
    except Exception as exc:
        print(f"[Scheduler] start failed: {exc}", flush=True)
        traceback.print_exc()
        return False


def _stop_scheduler_once():
    global _scheduler_started
    if not _scheduler_started:
        return
    try:
        _, stop_scheduler = _get_scheduler_funcs()
        if stop_scheduler is None:
            return
        stop_scheduler()
    except Exception:
        pass
    _scheduler_started = False


def _warmup_data_environment():
    # datahub/scheduler 初始化依赖 tdxpy，在 Windows 上频繁触发 NotImplementedError
    # 导致服务器崩溃。直接跳过，使用内置的 fallback 数据源（新浪/东方财富）。
    print("[Startup] DataHub warmup skipped, using fallback data sources", flush=True)
    return


def _fetch_fallback_news_payload(limit: int) -> dict:
    try:
        from backend.data_sources.news import news_source

        payload = news_source.get_news(limit=max(1, min(limit, 80)))
        if payload.get("items"):
            return payload
    except Exception:
        traceback.print_exc()

    items = []
    seen = set()
    errors = []
    sources = [
        ("https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid=2510&num=20&page=1", "新浪财经"),
        ("https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid=2513&num=20&page=1", "新浪证券"),
        ("https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid=2515&num=20&page=1", "新浪股票"),
        ("https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid=2518&num=20&page=1", "新浪滚动"),
    ]
    limit = max(1, min(int(limit or 40), 80))

    for target, source_name in sources:
        try:
            code, _, body = fetch_url(target)
            if code != 200 or not body:
                raise RuntimeError(f"HTTP {code}")
            payload = json.loads(body.decode("utf-8", errors="replace"))
            rows = (((payload or {}).get("result") or {}).get("data") or [])
            for row in rows:
                title = str(row.get("title") or "").strip()
                if not title:
                    continue
                url = str(row.get("url") or "").strip()
                key = (title, url)
                if key in seen:
                    continue
                seen.add(key)
                ts_raw = row.get("ctime")
                ts_val = int(ts_raw) if str(ts_raw or "").isdigit() else ts_raw
                items.append({
                    "id": row.get("oid") or row.get("docid") or f"{source_name}:{title}",
                    "title": title,
                    "summary": str(row.get("summary") or row.get("intro") or row.get("wapsummary") or title).strip(),
                    "url": url,
                    "source": source_name,
                    "time": ts_val,
                    "timeStr": "",
                    "sentiment": "neutral",
                    "tags": [],
                    "category": "important",
                    "importance": 10,
                    "freshness": "realtime",
                    "stale": False,
                    "errors": [],
                })
                if len(items) >= limit:
                    break
            if len(items) >= limit:
                break
        except Exception as exc:
            errors.append({"source": source_name, "message": str(exc)})

    return {
        "source": "新浪财经(兜底)",
        "updatedAt": now_text(),
        "asOf": now_text(),
        "freshness": "realtime" if items else "unavailable",
        "stale": not bool(items),
        "unavailable": not bool(items),
        "errors": errors,
        "data": {
            "items": items[:limit],
            "news": items[:limit],
            "count": len(items[:limit]),
        },
        "items": items[:limit],
        "news": items[:limit],
        "count": len(items[:limit]),
        "total": len(items[:limit]),
        "sources": ["新浪财经(兜底)"],
    }


def _normalize_symbol(symbol: str) -> str:
    symbol = str(symbol or "").strip().lower()
    if symbol.startswith(("sh", "sz", "bj")):
        return symbol
    if symbol.endswith(".sh"):
        return "sh" + symbol[:6]
    if symbol.endswith(".sz"):
        return "sz" + symbol[:6]
    if symbol.startswith(("6", "5", "9")):
        return "sh" + symbol[:6]
    if symbol.startswith(("0", "2", "3")):
        return "sz" + symbol[:6]
    return symbol


def _fetch_sina_kline(symbol: str, scale: int, length: int) -> list[dict]:
    sym = _normalize_symbol(symbol)
    # 尝试 HTTP 和 HTTPS 两个版本（HTTPS 常触发456反爬）
    for proto in ("http", "https"):
        try:
            url = (
                f"{proto}://money.finance.sina.com.cn/quotes_service/api/json_v2.php/"
                "CN_MarketData.getKLineData?"
                f"symbol={sym}&scale={int(scale)}&datalen={min(int(length), 800)}"
            )
            code, _, body = fetch_url(url)
            if code != 200 or not body:
                continue
            data = json.loads(body.decode("utf-8", errors="ignore"))
            if not isinstance(data, list):
                continue
            rows = []
            for item in data:
                if not isinstance(item, dict):
                    continue
                date = item.get("day") or item.get("date") or item.get("time") or ""
                if not date:
                    continue
                try:
                    rows.append({
                        "date": date,
                        "open": float(item.get("open", 0) or 0),
                        "close": float(item.get("close", 0) or 0),
                        "high": float(item.get("high", 0) or 0),
                        "low": float(item.get("low", 0) or 0),
                        "volume": float(item.get("volume", 0) or 0),
                        "amount": float(item.get("amount", 0) or 0),
                        "source": "sina",
                    })
                except Exception:
                    continue
            if rows:
                return rows[-length:]
        except Exception:
            continue
    return []


def _fetch_eastmoney_kline(symbol: str, scale: int, length: int) -> list[dict]:
    sym = _normalize_symbol(symbol)
    secid = f"1.{sym[2:]}" if sym.startswith("sh") else f"0.{sym[2:]}"
    if scale == 1:
        klt = 1
    elif scale >= 7200:
        klt = 103
    elif scale >= 1200:
        klt = 102
    elif scale >= 240:
        klt = 101
    elif scale >= 60:
        klt = 60
    elif scale >= 30:
        klt = 30
    elif scale >= 15:
        klt = 15
    else:
        klt = 5
    try:
        url = (
            "https://push2his.eastmoney.com/api/qt/stock/kline/get?"
            f"secid={secid}&klt={klt}&fqt=1&lmt={min(int(length), 800)}&end=20500101&smplmt=460"
            "&fields1=f1,f2,f3,f4,f5,f6"
            "&fields2=f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61"
        )
        req = urllib.request.Request(url, headers={**HEADERS, "Referer": "https://data.eastmoney.com/"})
        with _NO_PROXY_OPENER.open(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        klines = (data.get("data") or {}).get("klines") or []
        rows = []
        for row in klines:
            parts = row.split(",")
            if len(parts) < 6:
                continue
            try:
                rows.append({
                    "date": parts[0],
                    "open": float(parts[1] or 0),
                    "close": float(parts[2] or 0),
                    "high": float(parts[3] or 0),
                    "low": float(parts[4] or 0),
                    "volume": float(parts[5] or 0),
                    "amount": float(parts[6] or 0) if len(parts) > 6 else 0,
                    "source": "eastmoney",
                })
            except Exception:
                continue
        return rows[-length:]
    except Exception:
        return []


def _fetch_akshare_kline(symbol: str, scale: int, length: int) -> list[dict]:
    """AKShare K线数据源（新浪底层，绕过反爬）"""
    try:
        import akshare as ak
        sym = _normalize_symbol(symbol)
        is_index = sym.startswith("sh000") or sym.startswith("sz399") or sym.startswith("sh0003") or sym.startswith("sh0006")
        if scale >= 7200:
            period = "monthly"
        elif scale >= 1200:
            period = "weekly"
        elif scale >= 240:
            period = "daily"
        else:
            # AKShare 不支持分钟线直接获取，返回空
            return []
        if is_index:
            df = ak.stock_zh_index_daily(symbol=sym)
            rows = []
            for _, row in df.tail(min(length, 800)).iterrows():
                rows.append({
                    "date": str(row["date"])[:10],
                    "open": float(row["open"]),
                    "close": float(row["close"]),
                    "high": float(row["high"]),
                    "low": float(row["low"]),
                    "volume": float(row["volume"]),
                    "amount": 0,
                    "source": "akshare",
                })
            return rows[-length:]
        else:
            from datetime import datetime
            end_date = datetime.now().strftime("%Y%m%d")
            start_date_dt = datetime.now()
            from datetime import timedelta
            start_date = (start_date_dt - timedelta(days=length * 2)).strftime("%Y%m%d")
            df = ak.stock_zh_a_daily(symbol=sym, start_date=start_date, end_date=end_date, adjust="qfq")
            if period == "weekly":
                # 从日线聚合周线
                from collections import defaultdict
                weeks = defaultdict(list)
                for _, row in df.iterrows():
                    d = str(row["date"])[:10]
                    dt_obj = datetime.strptime(d, "%Y-%m-%d")
                    week_start = dt_obj - timedelta(days=dt_obj.weekday())
                    weeks[week_start.strftime("%Y-%m-%d")].append(row)
                rows = []
                for week_date, items in sorted(weeks.items()):
                    rows.append({
                        "date": week_date,
                        "open": float(items[0]["open"]),
                        "close": float(items[-1]["close"]),
                        "high": max(float(r["high"]) for r in items),
                        "low": min(float(r["low"]) for r in items),
                        "volume": sum(float(r["volume"]) for r in items),
                        "amount": sum(float(r.get("amount", 0)) for r in items),
                        "source": "akshare-weekly",
                    })
                return rows[-length:]
            rows = []
            for _, row in df.tail(min(length, 800)).iterrows():
                rows.append({
                    "date": str(row["date"])[:10],
                    "open": float(row["open"]),
                    "close": float(row["close"]),
                    "high": float(row["high"]),
                    "low": float(row["low"]),
                    "volume": float(row["volume"]),
                    "amount": float(row.get("amount", 0)),
                    "source": "akshare",
                })
            return rows[-length:]
    except Exception as exc:
        print(f"[AKShare Kline] {symbol} failed: {exc}", flush=True)
        return []


def _fetch_qq_kline(symbol: str, scale: int, length: int) -> list[dict]:
    """腾讯财经K线API（直连可用，作为主力K线数据源）"""
    sym = _normalize_symbol(symbol)
    # 映射 scale → 腾讯周期名
    if scale >= 7200:
        ktype = "month"
    elif scale >= 1200:
        ktype = "week"
    elif scale >= 240:
        ktype = "day"
    elif scale >= 60:
        ktype = "60"
    elif scale >= 30:
        ktype = "30"
    elif scale >= 15:
        ktype = "15"
    elif scale >= 5:
        ktype = "5"
    else:
        ktype = "1"
    try:
        url = (
            f"http://web.ifzq.gtimg.cn/appstock/app/fqkline/get?"
            f"param={sym},{ktype},,,{min(int(length), 800)},qfq"
        )
        req = urllib.request.Request(url, headers=HEADERS)
        with _NO_PROXY_OPENER.open(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        stock_data = (data.get("data") or {}).get(sym) or {}
        # 日线在 "day" 或 "qfqday" 字段
        rows_raw = stock_data.get("qfqday") or stock_data.get("day") or []
        if not rows_raw:
            # 周线/月线/分钟线
            for key in ("qfqweek", "week", "qfqmonth", "month", "m5", "m15", "m30", "m60"):
                if key in stock_data:
                    rows_raw = stock_data[key]
                    break
        rows = []
        for item in rows_raw:
            if not isinstance(item, (list, tuple)) or len(item) < 6:
                continue
            try:
                rows.append({
                    "date": str(item[0]),
                    "open": float(item[1] or 0),
                    "close": float(item[2] or 0),
                    "high": float(item[3] or 0),
                    "low": float(item[4] or 0),
                    "volume": float(item[5] or 0),
                    "amount": 0,
                    "source": "qq",
                })
            except Exception:
                continue
        return rows[-length:]
    except Exception:
        return []


def _fetch_eastmoney_trends_1m(symbol: str, length: int) -> list[dict]:
    sym = _normalize_symbol(symbol)
    secid = f"1.{sym[2:]}" if sym.startswith("sh") else f"0.{sym[2:]}"
    try:
        url = (
            "https://push2.eastmoney.com/api/qt/stock/trends2/get?"
            f"secid={secid}&fields1=f1,f2,f3,f4,f5,f6,f7,f8,f9,f10,f11"
            "&fields2=f51,f52,f53,f54,f55,f56,f57,f58&iscr=0&ndays=1"
        )
        req = urllib.request.Request(url, headers={**HEADERS, "Referer": "https://quote.eastmoney.com/"})
        with _NO_PROXY_OPENER.open(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        rows = []
        for row in ((data.get("data") or {}).get("trends") or []):
            parts = row.split(",")
            if len(parts) < 7:
                continue
            try:
                rows.append({
                    "date": parts[0],
                    "open": float(parts[1] or 0),
                    "close": float(parts[2] or 0),
                    "high": float(parts[3] or 0),
                    "low": float(parts[4] or 0),
                    "volume": float(parts[5] or 0),
                    "amount": float(parts[6] or 0),
                    "source": "eastmoney-trends",
                })
            except Exception:
                continue
        return rows[-length:]
    except Exception:
        return []


def _fetch_akshare_global_kline(symbol: str, length: int) -> list[dict]:
    """AKShare 全球指数 K 线兜底（如 KS11）"""
    try:
        from backend.services.market_service import market_service
        rows = market_service._fetch_global_index_kline(symbol, days=length)
        if rows:
            return rows
    except Exception:
        pass
    return []


def _get_kline_data(symbol: str, scale: int, length: int) -> list[dict]:
    symbol = str(symbol or "").strip().upper()
    # 全球指数优先走 AKShare
    if symbol in {"KS11", "N225", "HSI", "HSTECH", "GDAXI", "VIX"}:
        rows = _fetch_akshare_global_kline(symbol, length)
        if rows:
            return rows
    if int(scale) == 1:
        rows = _fetch_eastmoney_trends_1m(symbol, length)
        if rows:
            return rows
    rows = _fetch_qq_kline(symbol, scale, length)
    if rows:
        return rows
    rows = _fetch_sina_kline(symbol, scale, length)
    if rows:
        return rows
    rows = _fetch_eastmoney_kline(symbol, scale, length)
    if rows:
        return rows
    try:
        from backend.datahub import datahub_service
        return datahub_service.get_kline(symbol, scale, length)
    except Exception:
        return []


PORT = int(os.getenv("PORT", "8080"))
BASE_DIR = Path(__file__).resolve().parent
ALLOWED_HOSTS = {
    "hq.sinajs.cn",
    "money.finance.sina.com.cn",
    "feed.mix.sina.com.cn",
    "finance.sina.com.cn",
    "push2.eastmoney.com",
    "push2his.eastmoney.com",
    "data.eastmoney.com",
    "quote.eastmoney.com",
    "datacenter-web.eastmoney.com",
    "news.10jqka.com.cn",
    "data.10jqka.com.cn",
    "www.jin10.com",
    "web.ifzq.gtimg.cn",
    "qt.gtimg.cn",
}

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36"
    ),
    "Referer": "https://finance.sina.com.cn/",
    "Accept": "*/*",
}


def _clean_nan(obj):
    import math
    if isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return None
        return obj
    if isinstance(obj, list):
        return [_clean_nan(v) for v in obj]
    if isinstance(obj, dict):
        return {k: _clean_nan(v) for k, v in obj.items()}
    return obj

def _json_default(obj):
    from datetime import date, datetime
    import math
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()
    if isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return None
        return obj
    # numpy types
    try:
        import numpy as np
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.floating):
            if np.isnan(obj) or np.isinf(obj):
                return None
            return float(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, np.bool_):
            return bool(obj)
    except Exception:
        pass
    # pandas NaT
    try:
        import pandas as pd
        if obj is pd.NaT:
            return None
    except Exception:
        pass
    raise TypeError(f"Object of type {type(obj).__name__} is not JSON serializable")

def now_text() -> str:
    from datetime import datetime

    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


_KLINE_CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "kline_cache")


def _save_kline_cache(symbol, scale, data):
    try:
        os.makedirs(_KLINE_CACHE_DIR, exist_ok=True)
        fname = f"{symbol}_{scale}.json"
        path = os.path.join(_KLINE_CACHE_DIR, fname)
        payload = {"data": data, "symbol": symbol, "scale": scale, "source": data[0].get("source", "?") if data else "?", "cached_at": now_text()}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
    except Exception:
        pass


def _load_kline_cache(symbol, scale):
    try:
        fname = f"{symbol}_{scale}.json"
        path = os.path.join(_KLINE_CACHE_DIR, fname)
        if not os.path.exists(path):
            return None
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _is_retryable(exc: Exception) -> bool:
    """判断是否为可重试的网络错误"""
    msg = str(exc).lower()
    retryable = [
        "connection reset", "connection aborted", "broken pipe",
        "timed out", "timeout", "temporary failure", "try again",
        "no route to host", "connection refused",
    ]
    return any(kw in msg for kw in retryable)


# 绕过系统代理直连（避免 127.0.0.1:7897 代理不可用导致连接失败）
import ssl as _ssl
_ssl_ctx = _ssl.create_default_context()
_NO_PROXY_OPENER = urllib.request.build_opener(
    urllib.request.ProxyHandler({}),
    urllib.request.HTTPSHandler(context=_ssl_ctx)
)

def fetch_url(target_url: str, retries: int = 2) -> tuple[int, str, bytes]:
    host = urllib.parse.urlparse(target_url).hostname
    if host not in ALLOWED_HOSTS:
        return 403, "text/plain; charset=utf-8", f"Forbidden: {host}".encode("utf-8")
    last_error = None
    for attempt in range(retries + 1):
        try:
            req = urllib.request.Request(target_url, headers=HEADERS)
            with _NO_PROXY_OPENER.open(req, timeout=15) as resp:
                body = resp.read()
                if resp.headers.get("Content-Encoding", "") == "gzip":
                    body = gzip.decompress(body)
                content_type = resp.headers.get("Content-Type", "text/plain")
                return 200, content_type, body
        except urllib.error.HTTPError as exc:
            last_error = exc
            if attempt < retries:
                _time.sleep(1)
            continue
        except (ssl.SSLError, ssl.SSLEOFError) as exc:
            last_error = exc
            if attempt < retries:
                _time.sleep(1)
            continue
        except Exception as exc:
            last_error = exc
            if attempt < retries and _is_retryable(exc):
                _time.sleep(1)
                continue
            break
    if last_error is None:
        return 502, "text/plain; charset=utf-8", b"Unknown error"
    if isinstance(last_error, urllib.error.HTTPError):
        return last_error.code, "text/plain; charset=utf-8", str(last_error).encode("utf-8")
    return 502, "text/plain; charset=utf-8", str(last_error).encode("utf-8")


def export_excel(records: dict) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "review"
    ws.append(["date", records.get("date", "")])
    ws.append(["generated_at", now_text()])
    ws.append(["note", records.get("note", "")])
    stream = io.BytesIO()
    wb.save(stream)
    return stream.getvalue()


def export_word(records: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Daily Review", 0)
    doc.add_paragraph(f"date: {records.get('date', '')}")
    doc.add_paragraph(f"generated_at: {now_text()}")
    doc.add_paragraph(records.get("note", "") or "")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


class Handler(http.server.BaseHTTPRequestHandler):
    def _serve_datahub_or_fallback(self, path, params, handler):
        try:
            dh = _get_datahub()
            if dh is None:
                raise RuntimeError("datahub unavailable")
            payload = handler(dh)
            self._json(payload)
            return
        except Exception:
            traceback.print_exc()
            self._fallback_api(path, params)

    def do_GET(self):
        try:
            parsed = urllib.parse.urlparse(self.path)
            params = urllib.parse.parse_qs(parsed.query)
            path = parsed.path

            # 静态文件优先（不依赖 datahub）
            if path not in ("/api/health", "/api/portfolio", "/api/indices",
                           "/api/quote", "/api/news", "/api/kline",
                           "/api/fund-flow", "/api/datahub/latest",
                           "/api/data/snapshots", "/api/data/portfolio",
                           "/api/datahub/news", "/api/datahub/quotes",
                           "/api/datahub/portfolio", "/api/datahub/sectors",
                           "/api/datahub/sector-history", "/api/datahub/kline",
                           "/api/fundamental/indicators", "/api/fundamental/quarterly",
                           "/api/fundamental/balance", "/api/fundamental/cashflow",
                           "/api/fundamental/forecast", "/api/announcements/stock",
                           "/api/fundamental/comprehensive"):
                # 非数据API走原有逻辑
                pass
            else:
                # 数据API：跳过 datahub（tdxpy 在 Windows 上不稳定），
                # 直接使用内置 fallback 数据源（新浪/东方财富）。
                self._fallback_api(path, params)
                return

            if path == "/api/health":
                self._json({"status": "ok", "mode": "fallback", "datahub": False})
                return
            if path in {"/api/datahub/latest", "/api/data/latest"}:
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.latest())
            if path in {"/api/datahub/snapshots", "/api/data/snapshots"}:
                kind = params.get("kind", [None])[0]
                def _snap(dh):
                    items = dh.snapshot_list(kind)
                    return {
                        "source": "datahub_snapshots",
                        "updatedAt": dh.state().get("last_update_time"),
                        "asOf": "2026-06-23",
                        "freshness": "local-cache",
                        "stale": False,
                        "unavailable": False,
                        "errors": [],
                        "data": {"items": items},
                        "items": items,
                    }
                return self._serve_datahub_or_fallback(path, params, _snap)
            if path in ("/api/datahub/portfolio", "/api/portfolio"):
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.portfolio())
            if path in ("/api/datahub/quotes", "/api/quote", "/api/indices"):
                symbols = [s.strip() for s in ",".join(params.get("symbols", [])).split(",") if s.strip()]
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.quotes(symbols))
            if path == "/api/datahub/sectors":
                refresh = params.get("refresh", ["0"])[0] in {"1", "true", "yes"}
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.sectors(refresh=refresh))
            if path in ("/api/datahub/news", "/api/news"):
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.news())
            if path in ("/api/fund-flow", "/api/sector-flow", "/api/sector-flow-sina"):
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.sectors(refresh=False))
            if path == "/api/datahub/sector-history":
                name = params.get("name", [""])[0]
                days = int(params.get("days", ["30"])[0] or 30)
                return self._serve_datahub_or_fallback(path, params, lambda dh: dh.sector_history(name, days))
            if path == "/api/fundamental/indicators":
                self._json(self._fetch_fundamental_indicators(params.get("symbol", [""])[0]))
                return
            if path == "/api/fundamental/quarterly":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "利润表"))
                return
            if path == "/api/fundamental/balance":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "资产负债表"))
                return
            if path == "/api/fundamental/cashflow":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "现金流量表"))
                return
            if path == "/api/fundamental/forecast":
                self._json(self._fetch_earnings_forecast(params.get("symbol", [""])[0]))
                return
            if path == "/api/announcements/stock":
                self._json(self._fetch_stock_announcements(
                    params.get("symbol", [""])[0],
                    int(params.get("page", ["1"])[0] or 1),
                    int(params.get("pagesize", ["20"])[0] or 20)
                ))
                return
            if path == "/api/fundamental/comprehensive":
                self._json(self._fetch_comprehensive_fundamental(params.get("symbol", [""])[0]))
                return
            if path == "/api/proxy":
                self._proxy()
                return
            if path == "/api/fundflow/stock":
                self._json(self._fetch_stock_fundflow(params))
                return

            self._static()
        except Exception as _exc:
            _tb = traceback.format_exc()
            traceback.print_exc()
            self._send(500, "text/plain; charset=utf-8", ("Internal server error: " + str(_exc) + "\n" + _tb).encode("utf-8", errors="ignore"))

    def do_POST(self):
        try:
            parsed = urllib.parse.urlparse(self.path)
            path = parsed.path
            if path in {"/api/datahub/refresh", "/api/data/refresh"}:
                try:
                    dh = _get_datahub()
                    if dh is None:
                        raise RuntimeError("datahub unavailable")
                    self._json(dh.refresh(force=True))
                except Exception:
                    traceback.print_exc()
                    self._json({"status": "ok", "mode": "sina-fallback", "datahub": False, "message": "refresh skipped"})
                return
            if path == "/api/export/review":
                self._export_review()
                return
            self._send(404, "text/plain; charset=utf-8", b"Not found")
        except Exception:
            traceback.print_exc()
            self._send(500, "text/plain; charset=utf-8", b"Internal server error")

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "*")
        self.end_headers()

    def _proxy(self):
        params = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
        target_url = params.get("url", [None])[0]
        if not target_url:
            self._send(400, "text/plain; charset=utf-8", b"Missing url")
            return
        code, content_type, body = fetch_url(target_url)
        self._send(code, content_type, body)

    def _export_review(self):
        length = int(self.headers.get("Content-Length", "0"))
        payload = self.rfile.read(length).decode("utf-8")
        records = json.loads(payload or "{}")
        fmt = (records.get("format") or "xlsx").lower()
        if fmt == "docx":
            body = export_word(records)
            ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"daily_review_{records.get('date', '')}.docx"
        else:
            body = export_excel(records)
            ctype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            filename = f"daily_review_{records.get('date', '')}.xlsx"
        self.send_response(200)
        self.send_header("Content-Type", ctype)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", str(len(body)))
        try:
            self.end_headers()
            self.wfile.write(body)
        except (BrokenPipeError, ConnectionResetError, OSError):
            pass

    def _static(self):
        path = urllib.parse.unquote(self.path.split("?")[0])
        if path == "/":
            path = "/terminal.html"
        filepath = (BASE_DIR / path.lstrip("/")).resolve()
        if not str(filepath).startswith(str(BASE_DIR)) or not filepath.is_file():
            self._send(404, "text/plain; charset=utf-8", b"Not found")
            return
        content_type = mimetypes.guess_type(str(filepath))[0] or "application/octet-stream"
        if filepath.suffix == ".js":
            content_type = "application/javascript; charset=utf-8"
        elif filepath.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        self._send(200, content_type, filepath.read_bytes())

    def _json(self, payload):
        self._send(200, "application/json; charset=utf-8", json.dumps(_clean_nan(payload), ensure_ascii=False, default=_json_default).encode("utf-8"))

    def _send(self, code, content_type, body):
        try:
            self.send_response(code)
            self.send_header("Content-Type", content_type)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
        except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError, OSError):
            pass

    def log_message(self, format, *args):
        pass

    # ============================================================
    # Fallback：datahub 不可用时直接调新浪 API
    # ============================================================
    _SYMBOLS = ["sh600487","sz002475","sz002384","sz000988","sh600459",
                "sh603211","sh600206","sz000636"]
    _INDICES = ["sh000001","sz399001","sz399006","sh000300","sh000688","KS11"]

    _GLOBAL_INDEX_CODES = {"KS11", "N225", "HSI", "HSTECH", "GDAXI", "VIX"}

    def _fetch_stock_fundflow(self, params):
        """东财个股真实资金流（datacenter RPT_DMSK_TS_FUNDFLOWHIS）。

        直接请求东财数据中心接口，绕过系统代理，避免 push2his 连接不稳定问题。
        返回字段保持与前端兼容：mainNetInflow / midNetInflow / smallNetInflow。
        """
        symbol = params.get("symbol", [""])[0]
        days = int(params.get("days", ["5"])[0] or 5)
        if not symbol:
            return {"items": [], "source": "东财个股资金流", "error": "missing symbol"}

        sym = _normalize_symbol(symbol)
        sym6 = sym[2:]
        filter_str = f'(SECURITY_CODE=%22{sym6}%22)'
        url = (
            "https://datacenter.eastmoney.com/securities/api/data/get?"
            "type=RPT_DMSK_TS_FUNDFLOWHIS&sty=ALL&source=SECURITIES&client=APP&"
            "st=TRADE_DATE&sr=-1&p=1&"
            f"ps={max(1, min(int(days), 120))}&filter={filter_str}"
            "&extraCols=f66%7C02%7CSECURITY_CODE%7CSUPERDEAL_NET,"
            "f72%7C02%7CSECURITY_CODE%7CBIGDEAL_NET,"
            "f78%7C02%7CSECURITY_CODE%7CMIDDEAL_NET,"
            "f84%7C02%7CSECURITY_CODE%7CSMALLDEAL_NET"
            f"&_={int(_time.time() * 1000)}"
        )
        try:
            req = urllib.request.Request(url, headers={**HEADERS, "Referer": "https://data.eastmoney.com/"})
            with _NO_PROXY_OPENER.open(req, timeout=20) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            rows = (data.get("result") or {}).get("data") or []
            items = []
            for row in rows:
                trade_date = str(row.get("TRADE_DATE") or "")[:10]
                if not trade_date:
                    continue
                close_price = float(row.get("CLOSE_PRICE") or 0) or None
                change_rate = float(row.get("CHANGE_RATE") or 0) or None
                net_inflow = float(row.get("NET_INFLOW") or 0)
                super_net = float(row.get("SUPERDEAL_NET") or 0)
                big_net = float(row.get("BIGDEAL_NET") or 0)
                mid_net = float(row.get("MIDDEAL_NET") or 0)
                small_net = float(row.get("SMALLDEAL_NET") or 0)
                main_net = super_net + big_net
                retail_net = mid_net + small_net
                # 若拆分字段与 NET_INFLOW 不匹配（历史行会重复最新日拆分），
                # 以 NET_INFLOW 作为真实主力净流入，并按比例估算散户流向。
                if abs(main_net - net_inflow) > max(abs(net_inflow), 1) * 0.05:
                    main_net = net_inflow
                    if retail_net != 0:
                        mid_ratio = mid_net / retail_net if retail_net != 0 else 0.5
                        small_ratio = 1 - mid_ratio
                        retail_net = -net_inflow
                        mid_net = retail_net * mid_ratio
                        small_net = retail_net * small_ratio
                    else:
                        mid_net = -net_inflow / 2
                        small_net = -net_inflow / 2
                items.append({
                    "date": trade_date,
                    "close": close_price,
                    "change": change_rate,
                    "mainNetInflow": main_net,
                    "superLargeNetInflow": super_net,
                    "largeNetInflow": big_net,
                    "midNetInflow": mid_net,
                    "smallNetInflow": small_net,
                    "totalNetInflow": net_inflow,
                    "source": "东财个股资金流",
                })
            return {
                "items": items,
                "symbol": symbol,
                "days": days,
                "source": "东财个股资金流",
                "updatedAt": now_text(),
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"items": [], "symbol": symbol, "source": "东财个股资金流", "error": str(e)}

    def _fetch_global_index_fallback(self, symbols):
        """datahub 不可用时，通过 market_service 获取全球指数"""
        result = {}
        codes = [s for s in symbols if str(s).strip().upper() in self._GLOBAL_INDEX_CODES]
        if not codes:
            return result
        try:
            from backend.services.market_service import market_service
            overview = market_service.get_index_overview()
            for item in overview.get("global", []):
                if item.get("code") in codes and item.get("close") is not None:
                    code = item["code"]
                    result[code] = {
                        "name": item.get("name", code),
                        "current": item.get("close"),
                        "prevClose": item.get("prev_close"),
                        "open": item.get("open"),
                        "high": item.get("high"),
                        "low": item.get("low"),
                        "volume": item.get("volume", 0),
                        "amount": item.get("amount", 0),
                        "change": item.get("change_pct"),
                        "pct": item.get("change_pct"),
                        "source": item.get("source", "akshare"),
                    }
        except Exception:
            pass
        return result

    def _fetch_sina(self, symbols):
        """从新浪获取实时行情"""
        url = "http://hq.sinajs.cn/list=" + ",".join(symbols)
        code, ct, body = fetch_url(url)
        if code != 200:
            return {}
        result = {}
        for line in body.decode("gbk", errors="replace").split(";"):
            line = line.strip()
            if not line or "var hq_str_" not in line:
                continue
            parts = line.split('"')
            if len(parts) < 2:
                continue
            sym = parts[0].replace("var hq_str_", "")
            vals = [v.strip() for v in parts[1].split(",")]
            if len(vals) < 32:
                continue
            result[sym] = {
                "name": vals[0], "current": vals[3],
                "open": vals[2], "high": vals[4], "low": vals[5],
                "volume": vals[8], "amount": vals[9],
                "change": vals[2] and vals[3] and float(vals[3]) - float(vals[2]) if vals[2] and vals[3] else 0,
                "pct": (float(vals[3]) - float(vals[2])) / float(vals[2]) * 100 if vals[2] and vals[3] and float(vals[2]) else 0,
                "turnover": vals[38] if len(vals) > 38 else "",
                "volume_ratio": "", "source": "新浪财经"
            }
        return result

    # ============================================================
    # 财务数据抓取（AKShare）
    # ============================================================
    def _symbol_to_pure(self, symbol: str) -> str:
        s = str(symbol or "").strip().lower()
        return s.replace("sh", "").replace("sz", "").replace("bj", "")

    def _fetch_fundamental_indicators(self, symbol: str):
        """AKShare 主要财务指标"""
        code = self._symbol_to_pure(symbol)
        try:
            import akshare as ak
            df = ak.stock_financial_analysis_indicator(symbol=code)
            if df is None or df.empty:
                return {"items": [], "source": "akshare"}
            records = df.head(8).to_dict(orient="records")
            return {"items": records, "source": "akshare", "symbol": symbol, "count": len(records)}
        except Exception as e:
            return {"items": [], "source": "akshare", "symbol": symbol, "error": str(e)}

    def _fetch_fundamental_report(self, symbol: str, report_type: str):
        """AKShare 新浪财经财务报表（利润表/资产负债表/现金流量表）"""
        code = self._symbol_to_pure(symbol)
        try:
            import akshare as ak
            df = ak.stock_financial_report_sina(stock=code, symbol=report_type)
            if df is None or df.empty:
                return {"items": [], "source": "akshare", "type": report_type}
            records = df.head(12).to_dict(orient="records")
            return {"items": records, "source": "akshare", "symbol": symbol, "type": report_type, "count": len(records)}
        except Exception as e:
            return {"items": [], "source": "akshare", "symbol": symbol, "type": report_type, "error": str(e)}

    def _fetch_earnings_forecast(self, symbol: str):
        """AKShare 业绩预报 + 业绩快报 + 业绩报表"""
        code = self._symbol_to_pure(symbol)
        result = {"forecast": {"items": [], "source": "akshare"},
                  "express": {"items": [], "source": "akshare"},
                  "report": {"items": [], "source": "akshare"}}
        try:
            import akshare as ak
            # 业绩预报
            for date in ["20250630", "20241231", "20240630"]:
                try:
                    df = ak.stock_yjyg_em(date=date)
                    if df is not None and not df.empty:
                        mask = df["股票代码"].astype(str).str.strip() == code
                        if mask.any():
                            result["forecast"]["items"] = df[mask].head(5).to_dict(orient="records")
                            result["forecast"]["date"] = date
                            break
                except Exception:
                    continue
            # 业绩快报
            for date in ["20250630", "20241231", "20240630"]:
                try:
                    df = ak.stock_yjkb_em(date=date)
                    if df is not None and not df.empty:
                        mask = df["股票代码"].astype(str).str.strip() == code
                        if mask.any():
                            result["express"]["items"] = df[mask].head(5).to_dict(orient="records")
                            result["express"]["date"] = date
                            break
                except Exception:
                    continue
            # 业绩报表
            for date in ["20250630", "20241231", "20240630"]:
                try:
                    df = ak.stock_yjbb_em(date=date)
                    if df is not None and not df.empty:
                        mask = df["股票代码"].astype(str).str.strip() == code
                        if mask.any():
                            result["report"]["items"] = df[mask].head(5).to_dict(orient="records")
                            result["report"]["date"] = date
                            break
                except Exception:
                    continue
            result["symbol"] = symbol
            return result
        except Exception as e:
            result["error"] = str(e)
            return result

    def _fetch_stock_announcements(self, symbol: str, page: int = 1, pagesize: int = 20):
        """AKShare 个股公告"""
        code = self._symbol_to_pure(symbol)
        try:
            import akshare as ak
            df = ak.stock_notice_report(symbol=code)
            if df is None or df.empty:
                return {"items": [], "source": "akshare", "symbol": symbol}
            start = (page - 1) * pagesize
            end = start + pagesize
            records = df.iloc[start:end].to_dict(orient="records") if len(df) > start else []
            return {"items": records, "source": "akshare", "symbol": symbol, "count": len(records), "total": len(df)}
        except Exception as e:
            return {"items": [], "source": "akshare", "symbol": symbol, "error": str(e)}

    def _fetch_tonghuashun_forecast(self, symbol: str):
        """AKShare 同花顺业绩预测 + 机构评级"""
        code = self._symbol_to_pure(symbol)
        result = {"forecast": {"items": [], "source": "ths"},
                  "rating": {"items": [], "source": "ths"}}
        try:
            import akshare as ak
            # 同花顺盈利预测
            try:
                df = ak.stock_yjyg_ths(symbol=code)
                if df is not None and not df.empty:
                    result["forecast"]["items"] = df.head(10).to_dict(orient="records")
            except Exception:
                pass
            # 机构评级
            try:
                df2 = ak.stock_institute_recommend_detail(symbol=code)
                if df2 is not None and not df2.empty:
                    result["rating"]["items"] = df2.head(10).to_dict(orient="records")
            except Exception:
                pass
            result["symbol"] = symbol
            return result
        except Exception as e:
            result["error"] = str(e)
            return result

    def _fetch_eastmoney_detail(self, symbol: str):
        """东方财富个股详细财务摘要(web API兜底)"""
        code = self._symbol_to_pure(symbol)
        result = {"items": [], "source": "eastmoney-web"}
        try:
            url = (
                "https://datacenter-web.eastmoney.com/api/data/v1/get?"
                "sortColumns=UPDATE_DATE,SECURITY_CODE&sortTypes=-1,-1"
                "&pageSize=50&pageNumber=1&reportName=RPT_FCI_PERFORMANCEE"
                f'&columns=ALL&filter=(SECURITY_CODE="{code}")'
            )
            req = urllib.request.Request(url, headers={**HEADERS, "Referer": "https://data.eastmoney.com/"})
            with _NO_PROXY_OPENER.open(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            items = (data.get("result") or {}).get("data") or []
            result["items"] = items
            result["count"] = len(items)
        except Exception as e:
            result["error"] = str(e)
        return result

    def _fetch_sina_financial_abstract(self, symbol: str):
        """新浪财经财务摘要"""
        code = self._symbol_to_pure(symbol)
        result = {"items": [], "source": "sina"}
        try:
            import akshare as ak
            # 新浪财务指标摘要
            try:
                df = ak.stock_financial_analysis_indicator(symbol=code)
                if df is not None and not df.empty:
                    result["items"] = df.head(8).to_dict(orient="records")
            except Exception:
                pass
            result["symbol"] = symbol
        except Exception as e:
            result["error"] = str(e)
        return result

    def _fetch_comprehensive_fundamental(self, symbol: str):
        """综合多源财务数据: 东方财富 + 新浪财经 + 同花顺 + AKShare (并发采集)"""
        from concurrent.futures import ThreadPoolExecutor
        code = self._symbol_to_pure(symbol)
        result = {
            "symbol": symbol,
            "sources": [],
            "indicators": {"items": [], "source": ""},
            "quarterly": {"items": [], "source": ""},
            "balance": {"items": [], "source": ""},
            "cashflow": {"items": [], "source": ""},
            "forecast": {"items": [], "source": ""},
            "announcements": {"items": [], "source": ""},
            "eastmoney_detail": {"items": [], "source": ""},
            "ths_forecast": {"items": [], "source": ""},
        }

        tasks = {
            "indicators": (self._fetch_fundamental_indicators, (symbol,)),
            "quarterly": (self._fetch_fundamental_report, (symbol, "利润表")),
            "balance": (self._fetch_fundamental_report, (symbol, "资产负债表")),
            "cashflow": (self._fetch_fundamental_report, (symbol, "现金流量表")),
            "forecast": (self._fetch_earnings_forecast, (symbol,)),
            "announcements": (self._fetch_stock_announcements, (symbol, 1, 30)),
            "eastmoney_detail": (self._fetch_eastmoney_detail, (symbol,)),
            "ths_forecast": (self._fetch_tonghuashun_forecast, (symbol,)),
            "sina_abstract": (self._fetch_sina_financial_abstract, (symbol,)),
        }

        with ThreadPoolExecutor(max_workers=6) as exe:
            futures = {k: exe.submit(fn, *args) for k, (fn, args) in tasks.items()}
            for k, fut in futures.items():
                try:
                    result[k] = fut.result(timeout=25)
                except Exception:
                    result[k] = {"items": [], "source": k, "error": "timeout or failed"}

        source_map = {
            "indicators": "akshare-indicators",
            "quarterly": "akshare-quarterly",
            "balance": "akshare-balance",
            "cashflow": "akshare-cashflow",
            "forecast": "eastmoney-forecast",
            "announcements": "akshare-announcements",
            "eastmoney_detail": "eastmoney-detail",
            "ths_forecast": "tonghuashun",
            "sina_abstract": "sina",
        }
        for k, label in source_map.items():
            val = result[k]
            if k == "forecast":
                if val.get("forecast", {}).get("items") or val.get("express", {}).get("items"):
                    result["sources"].append(label)
            elif k == "ths_forecast":
                if val.get("forecast", {}).get("items") or val.get("rating", {}).get("items"):
                    result["sources"].append(label)
            elif val.get("items"):
                result["sources"].append(label)
        return result

    def _fallback_api(self, path, params):
        """datahub 不可用时的 fallback 数据接口"""
        try:
            if path == "/api/health":
                self._json({"status": "ok", "mode": "sina-fallback", "datahub": False})
                return
            if path in ("/api/portfolio", "/api/datahub/portfolio"):
                data = self._fetch_sina(self._SYMBOLS)
                self._json({"items": data, "count": len(data), "source": "新浪财经(fallback)"})
                return
            if path in ("/api/indices", "/api/quote", "/api/datahub/quotes"):
                syms = params.get("symbols", [])
                if not syms:
                    syms = self._INDICES + self._SYMBOLS
                data = self._fetch_sina(syms)
                data.update(self._fetch_global_index_fallback(syms))
                self._json(data)
                return
            if path == "/api/kline" or path == "/api/datahub/kline":
                symbol = params.get("symbol", ["sh000001"])[0]
                scale = int(params.get("scale", ["240"])[0] or 240)
                length = int(params.get("length", ["120"])[0] or 120)
                # 全球指数优先走 AKShare
                sym_upper = str(symbol).strip().upper()
                if sym_upper in self._GLOBAL_INDEX_CODES:
                    rows = _fetch_akshare_global_kline(sym_upper, length)
                    if rows:
                        self._json({"data": rows, "symbol": symbol, "scale": scale, "length": length, "source": "akshare-global"})
                        return
                # 优先AKShare K线（新浪底层，绕过反爬，日线/周线/月线稳定）
                rows = _fetch_akshare_kline(symbol, scale, length)
                if rows:
                    _save_kline_cache(symbol, scale, rows)
                    self._json({"data": rows, "symbol": symbol, "scale": scale, "length": length, "source": "akshare"})
                    return
                # 腾讯K线API（分钟线备用）
                rows = _fetch_qq_kline(symbol, scale, length)
                if rows:
                    _save_kline_cache(symbol, scale, rows)
                    self._json({"data": rows, "symbol": symbol, "scale": scale, "length": length, "source": "qq"})
                    return
                # 备用东方财富K线API
                rows = _fetch_eastmoney_kline(symbol, scale, length)
                if rows:
                    _save_kline_cache(symbol, scale, rows)
                    self._json({"data": rows, "symbol": symbol, "scale": scale, "length": length, "source": "eastmoney"})
                    return
                # 备用新浪K线API
                rows = _fetch_sina_kline(symbol, scale, length)
                if rows:
                    _save_kline_cache(symbol, scale, rows)
                    self._json({"data": rows, "symbol": symbol, "scale": scale, "length": length, "source": "sina"})
                    return
                # 所有外部API均失败，回退本地缓存
                cached = _load_kline_cache(symbol, scale)
                if cached:
                    self._json({"data": cached["data"], "symbol": symbol, "scale": scale, "length": length, "source": "cache", "cached_at": cached.get("cached_at", "")})
                    return
                self._json({"data": [], "error": "kline_unavailable"})
                return
            if path in ("/api/news", "/api/datahub/news"):
                limit = int(params.get("limit", ["40"])[0] or 40)
                self._json(_fetch_fallback_news_payload(limit))
                return
            # 财务数据 fallback
            if path == "/api/fundamental/indicators":
                self._json(self._fetch_fundamental_indicators(params.get("symbol", [""])[0]))
                return
            if path == "/api/fundamental/quarterly":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "利润表"))
                return
            if path == "/api/fundamental/balance":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "资产负债表"))
                return
            if path == "/api/fundamental/cashflow":
                self._json(self._fetch_fundamental_report(params.get("symbol", [""])[0], "现金流量表"))
                return
            if path == "/api/fundamental/forecast":
                self._json(self._fetch_earnings_forecast(params.get("symbol", [""])[0]))
                return
            if path == "/api/announcements/stock":
                self._json(self._fetch_stock_announcements(
                    params.get("symbol", [""])[0],
                    int(params.get("page", ["1"])[0] or 1),
                    int(params.get("pagesize", ["20"])[0] or 20)
                ))
                return
            if path == "/api/fundamental/comprehensive":
                self._json(self._fetch_comprehensive_fundamental(params.get("symbol", [""])[0]))
                return
            # 其他API返回空
            self._json({"items": {}, "count": 0, "source": "unavailable"})
        except Exception as e:
            self._json({"error": str(e), "items": {}})


if __name__ == "__main__":
    print(f"Server starting on http://127.0.0.1:{PORT}/terminal.html", flush=True)

    http.server.ThreadingHTTPServer.daemon_threads = True
    server = http.server.ThreadingHTTPServer(("127.0.0.1", PORT), Handler)
    server.socket.setsockopt(_socket.SOL_SOCKET, _socket.SO_REUSEADDR, 1)
    server.daemon_threads = True
    server.timeout = 1
    print("Server Running on port %d (Ctrl+C停止)" % PORT, flush=True)

    # 启动调度器（后台线程，延迟启动避免阻塞）
    def _delayed_scheduler_start():
        _time.sleep(5)
        _start_scheduler_once()

    threading.Thread(target=_delayed_scheduler_start, daemon=True).start()

    try:
        server.serve_forever(poll_interval=0.5)
    except KeyboardInterrupt:
        print("\nStopped", flush=True)
    except Exception as e:
        print("\n[Server Error] %s" % e, flush=True)
        traceback.print_exc()
    finally:
        print("[Server] shutting down...", flush=True)
        try:
            _stop_scheduler_once()
        except Exception:
            pass
        try:
            server.shutdown()
        except Exception:
            pass
        try:
            server.server_close()
        except Exception:
            pass
        print("[Server] stopped", flush=True)
